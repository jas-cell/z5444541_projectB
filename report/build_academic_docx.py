#!/usr/bin/env python3
"""Build a readable Word copy of the current academic report. Numbers locked to source."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image as PILImage

ROOT = Path(__file__).resolve().parent
FIGDIR = ROOT.parent / "results" / "figures"
OUT = ROOT / "report.docx"

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x44, 0x44, 0x44)
PALE = "F4F4F4"
LINE = "B0B0B0"


def set_run_font(run, *, name="Times New Roman", size=10.5, bold=False, italic=False, color=INK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, fill: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), LINE)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_keep_with_next(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    kwn = OxmlElement("w:keepNext")
    pPr.append(kwn)


def add_markup(paragraph, text: str, *, size=10.5, italic=False, color=INK, bold=False) -> None:
    """Render a small subset of reportlab markup: <i>, <b>, and plain text."""
    parts = re.split(r"(</?(?:i|b)>)", text)
    state_i = italic
    state_b = bold
    for part in parts:
        if part == "<i>":
            state_i = True
        elif part == "</i>":
            state_i = False
        elif part == "<b>":
            state_b = True
        elif part == "</b>":
            state_b = False
        elif part:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, bold=state_b, italic=state_i, color=color)


def add_para(
    doc,
    text: str,
    *,
    style="body",
    space_before=0,
    space_after=8,
    first_indent=0,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=10.5,
    bold=False,
    italic=False,
    color=INK,
    keep_with_next=False,
):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    if keep_with_next:
        set_keep_with_next(p)
    add_markup(p, text, size=size, italic=italic, color=color, bold=bold)
    return p


def h1(doc, text: str) -> None:
    add_para(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=16,
        space_after=8,
        size=13,
        bold=True,
        keep_with_next=True,
    )


def h2(doc, text: str) -> None:
    add_para(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=12,
        space_after=6,
        size=11.5,
        bold=True,
        keep_with_next=True,
    )


def body(doc, text: str, *, indent=True) -> None:
    add_para(doc, text, first_indent=0.45 if indent else 0, space_after=8, size=10.5)


def caption(doc, text: str) -> None:
    add_para(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=6,
        space_after=8,
        size=9,
        italic=True,
        keep_with_next=True,
    )


def note(doc, text: str) -> None:
    add_para(doc, text, space_before=4, space_after=10, size=9.5, italic=True, color=MUTED)


def eq(doc, text: str) -> None:
    add_para(
        doc,
        text,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=4,
        space_after=4,
        size=10.5,
    )


def eqnote(doc, text: str) -> None:
    add_para(doc, text, space_before=0, space_after=8, size=9.5, first_indent=0)


def ref(doc, text: str) -> None:
    p = add_para(doc, text, space_after=5, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)


def add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(h)
        set_run_font(run, size=8, bold=True)
        shade_cell(cell, PALE)
        set_cell_border(cell)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(val)
            set_run_font(run, size=8)
            set_cell_border(cell)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_figure(doc, filename: str, cap: str) -> None:
    src = FIGDIR / filename
    im = PILImage.open(src)
    max_w_in = 6.15
    max_h_in = 7.4
    native_w, native_h = im.size
    aspect = native_w / native_h
    width_in = max_w_in
    height_in = width_in / aspect
    if height_in > max_h_in:
        height_in = max_h_in
        width_in = height_in * aspect

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    set_keep_with_next(p)
    run = p.add_run()
    run.add_picture(str(src), width=Inches(width_in))
    add_para(
        doc,
        cap,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=4,
        space_after=10,
        size=9,
        italic=True,
    )


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.7)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left = hp.add_run("Kaiyuan Lan (z5444541)  |  FINS3645 Project B")
    set_run_font(left, size=8, italic=True, color=INK)
    # Right-aligned product name via a right tab at the content width.
    hp.paragraph_format.tab_stops.add_tab_stop(Cm(16.6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
    hp.add_run("\t")
    right = hp.add_run("Signal Harbour")
    set_run_font(right, size=8, italic=True, color=INK)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # PAGE field
    run1 = fp.add_run()
    set_run_font(run1, size=9)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    run1._r.append(fld1)
    run2 = fp.add_run()
    set_run_font(run2, size=9)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._r.append(instr)
    run3 = fp.add_run()
    set_run_font(run3, size=9)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run3._r.append(fld2)

    core = doc.core_properties
    core.title = "Signal Harbour: FINS3645 Project B"
    core.author = "Kaiyuan Lan (z5444541)"
    core.subject = "FINS3645 Project B"


def build() -> Path:
    doc = Document()
    configure_document(doc)

    add_para(doc, "Signal Harbour", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, size=16, bold=True)
    add_para(
        doc,
        "Systematic multi-asset funds, a sector news-sentiment index, and a Streamlit research app",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
        size=11,
        italic=True,
    )
    add_para(doc, "FINS3645 Project B", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1, size=10, color=MUTED)
    add_para(doc, "Kaiyuan Lan (z5444541)", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1, size=10, color=MUTED)
    add_para(
        doc,
        "Sample through December 2023",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
        size=10,
        color=MUTED,
    )

    h1(doc, "1. Customer, data, and backtest design")
    body(
        doc,
        "Signal Harbour is a research product for a small wealth-advisory firm. The day-to-day user is a "
        "portfolio analyst who needs comparable walk-forward evidence before an investment-committee discussion. "
        "Equity, Crypto, and Combined sleeves make up the menu, each run with equal weight, minimum variance, "
        "maximum Sharpe, and risk parity (Markowitz, 1952). Hutto and Gilbert (2014) supply the VADER baseline "
        "the course builds on.",
        indent=False,
    )
    body(
        doc,
        "Prices are adjusted-close simple returns within ticker. Equity and Combined sleeves use a 252-session "
        "estimation window and 252-day annualisation, whereas crypto sleeves sit on a 365-day native window. "
        "For Combined funds, crypto returns are computed on that native calendar first, then left-aligned onto "
        "equity dates, so weekend compounding is neither invented nor deleted in the merge.",
    )
    body(
        doc,
        "On monthly decision date <i>T</i>, target weights use returns strictly before <i>T</i>, become effective "
        "on the next session, and leave the close-to-close return on <i>T</i> with the prior holdings. Between "
        "rebalances the book drifts with relative performance. Optimisation is long-only, a 25% name cap applies "
        "in the broad sleeves, and covariance is annualised inside the solver. A candidate solution is kept only "
        "when it improves on equal weight.",
    )
    h2(doc, "Backtest conventions")
    body(
        doc,
        "The risk-free rate is zero in the Maximum Sharpe objective and in every reported Sharpe ratio. Headline "
        "fund results assume zero trading costs, with Table 5 later applying the same turnover-cost stresses to "
        "the MinVar base and the sentiment-fusion sleeve. The Streamlit application reads precomputed CSV files "
        "and does not re-run VADER, rebuild the sentiment series, or execute backtests at runtime.",
    )
    h2(doc, "First-live timing")
    body(
        doc,
        "The decision date and the first earning date are not the same under this convention. After the "
        "252-session warm-up, the first Equity and Combined decision date is 4 January 2021 and those targets "
        "first earn a return on 5 January 2021, while on the native crypto calendar the first decision date is "
        "1 January 2021 and the first live return is 2 January 2021. Figure A1 starts at those live dates, not "
        "at the beginning of the raw price histories.",
    )
    h2(doc, "Why timing and drift are specified this way")
    body(
        doc,
        "A desk that forms weights after the close cannot book the full close-to-close return on <i>T</i> as if "
        "the new book had already been traded at the previous close. Assigning <i>T</i> to the old holdings and "
        "switching on the next session is the conservative sequence. It is also reproducible.",
    )
    body(
        doc,
        "Drift is a separate issue. Resetting to target weights every day is not monthly rebalancing, even if "
        "the optimiser only runs once a month. Letting prices move the holdings, then measuring turnover against "
        "those drifted weights, keeps the wealth path and the cost exercise aligned. Annualising covariance inside "
        "the optimiser, and rejecting solutions that fail to beat equal weight, also answers a brief-level "
        "problem: SLSQP can report success without leaving its starting vector when daily variances are small. "
        "In the retained Equity sample, Minimum Variance differs from Equal Weight at every monthly decision, "
        "and the mean half-L1 distance is about 0.72, which is enough to show the solver actually moved the "
        "weights.",
    )

    h1(doc, "2. Out-of-sample fund results")
    body(
        doc,
        "Table 1 and Figure A1 report the full fund menu. Among mixed-asset portfolios, Combined Risk Parity has "
        "the highest zero-risk-free-rate Sharpe ratio, 0.91. CAGR is 14.3%, volatility is 16.2%, and maximum "
        "drawdown is -19.3% (Figures A2 and A4). Combined Maximum Sharpe finishes with slightly more wealth, "
        "1.55x and a 15.7% CAGR, but its drawdown reaches -26.2%. Both paths are research paths under the stated "
        "assumptions.",
        indent=False,
    )
    h2(doc, "How to read the mixed-asset ranking")
    body(
        doc,
        "The mixed result is mainly about diversification under a calendar a desk can actually trade. Combined "
        "funds are different. A stand-alone crypto portfolio compounds through weekends because it is evaluated "
        "on 365 native days, whereas Combined funds align crypto returns to equity sessions so the mixed desk "
        "earns that leg only when the equity book can also be traded. Under that constraint, Risk Parity "
        "repeatedly cuts exposure to the most volatile names in the trailing covariance estimate, without "
        "forecasting expected returns.",
    )
    body(
        doc,
        "Maximum Sharpe has a heavier estimation burden. It has to estimate means as well as covariance, and "
        "over a three-year live window noisy means encourage larger bets, which shows up in the deeper "
        "drawdown and in the December 2023 snapshot in Table 3. For a small advisory firm the distinction "
        "is practical, with Combined Risk Parity the stronger default comparison and Combined Maximum Sharpe "
        "belonging in a satellite discussion with clients who accept concentration and drawdown risk.",
    )
    h2(doc, "Equity and crypto sleeves")
    body(
        doc,
        "On equities, Equal Weight leads with a Sharpe ratio of 0.86. Risk Parity reaches 0.76, ahead of Minimum "
        "Variance at 0.47, but it still does not displace the transparent baseline. That does not make "
        "optimisation pointless. Once the walk-forward rule is imposed, this sample simply did not pay enough "
        "for the extra estimation to beat Equal Weight, and Minimum Variance, which does not forecast expected "
        "returns, has a weaker realised Sharpe more plausibly linked to unstable covariance estimates and binding "
        "caps than to mean estimation.",
    )
    body(
        doc,
        "Crypto Minimum Variance is the strongest crypto sleeve, with a 1.04 Sharpe ratio and a 62.1% CAGR. "
        "Crypto Maximum Sharpe is the reason CAGR and Sharpe are kept separate in this report. CAGR is -2.7%. "
        "The annualised arithmetic mean-to-volatility ratio remains positive at 0.35.",
    )
    caption(
        doc,
        "Table 1. Out-of-sample performance. CAGR is the annualised return from the growth-of-one path. "
        "Sharpe is annualised arithmetic mean divided by annualised volatility, risk-free rate zero. "
        "Equity-calendar funds use 252 days; native crypto funds use 365. Source: performance_metrics.csv. "
        "See Figures A1 and A4.",
    )
    add_table(
        doc,
        ["Fund", "CAGR", "Vol.", "Sharpe", "MaxDD", "Terminal"],
        [
            ["Eq EW", "13.3%", "16.1%", "0.86", "-20.2%", "1.45x"],
            ["Eq MinVar", "5.4%", "12.7%", "0.47", "-14.9%", "1.17x"],
            ["Eq MaxSharpe", "7.0%", "17.5%", "0.48", "-22.8%", "1.22x"],
            ["Eq RP", "10.4%", "14.5%", "0.76", "-18.4%", "1.34x"],
            ["Cry EW", "37.1%", "81.7%", "0.80", "-81.7%", "2.58x"],
            ["Cry MinVar", "62.1%", "71.2%", "1.04", "-71.5%", "4.25x"],
            ["Cry MaxSharpe", "-2.7%", "76.2%", "0.35", "-85.6%", "0.92x"],
            ["Cry RP", "41.0%", "79.7%", "0.83", "-80.1%", "2.80x"],
            ["Comb EW", "15.5%", "21.5%", "0.78", "-27.7%", "1.54x"],
            ["Comb MinVar", "5.4%", "12.8%", "0.48", "-15.1%", "1.17x"],
            ["Comb MaxSharpe", "15.7%", "23.3%", "0.74", "-26.2%", "1.55x"],
            ["Comb RP", "14.3%", "16.2%", "0.91", "-19.3%", "1.49x"],
            ["Eq MinVar+Sent", "5.3%", "12.7%", "0.47", "-14.9%", "1.17x"],
        ],
    )

    h2(doc, "Concentration and the Combined sleeve")
    body(
        doc,
        "Figure A3 follows the decision weights of Combined Minimum Variance and Risk Parity, with an Other "
        "residual so that each stacked area sums to one. Risk Parity stays more dispersed. That matches the "
        "shallower drawdown in Figure A2. Minimum Variance is willing to cluster in names that look "
        "low-volatility in the trailing covariance matrix. The clustering is coherent with its objective. It "
        "still becomes a problem when the selected cluster later weakens.",
    )
    body(
        doc,
        "Across families, the Combined sleeve tempers crypto-native volatility while keeping exposure to both "
        "asset classes. Crypto Equal Weight and Crypto Risk Parity reach high terminal growth, but their "
        "drawdowns remain extreme. The -2.7% CAGR and positive 0.35 Sharpe for Crypto Maximum Sharpe also warn "
        "against reading one metric in isolation. Once the mixed fund is evaluated on the equity calendar, "
        "Combined Risk Parity is the more usable default for a mixed-client conversation. The best pure-crypto "
        "result is not.",
    )
    caption(
        doc,
        "Table 2. Latest Combined Risk Parity decision weights, top five. Source: latest_holdings_snapshot.csv.",
    )
    add_table(
        doc,
        ["Ticker", "Weight"],
        [["MRK", "3.81%"], ["ABBV", "3.74%"], ["WMT", "3.37%"], ["KO", "3.19%"], ["TMUS", "3.17%"]],
    )
    caption(
        doc,
        "Table 3. Latest Combined Maximum Sharpe decision weights, top five. Source: latest_holdings_snapshot.csv.",
    )
    add_table(
        doc,
        ["Ticker", "Weight"],
        [["GE", "25.00%"], ["NVDA", "20.09%"], ["SO", "16.86%"], ["ADBE", "10.32%"], ["BTC-USD", "10.29%"]],
    )
    note(
        doc,
        "These tables are end-of-sample snapshots. They describe concentration at the final decision date, not "
        "the full path from 2021 to 2023. A large December 2023 weight cannot explain performance earlier in the "
        "sample.",
    )

    h1(doc, "3. Sentiment model, sector index, and fusion")
    body(
        doc,
        "Part B keeps the headline-processing rules from Part A. Ticker-date-title duplicates are removed, and "
        "each headline is mapped to the same or next equity session before Base VADER and the Week 9 finVADER "
        "analyser are re-scored without alteration for comparison. The Week 9 benchmark combines NLTK VADER, "
        "SentiBigNomics scaled by 0.1, and the Henry lexicon, matching the course helper and Korab’s FinVADER "
        "resources. Loughran and McDonald (2011) give the wider rationale for finance-specific dictionaries.",
        indent=False,
    )
    body(
        doc,
        "Signal Harbour starts from that Week 9 base and adds documented masks, finance terms, and "
        "negation-aware phrases. The resulting score is multiplied by coverage confidence and by the elevated "
        "Attention Pulse confidence developed in Part A. Abnormally low headline volume never raises confidence. "
        "Missing ticker-days are inserted as neutral on a complete equity ticker-day grid before sector "
        "aggregation, so Figure A5 and the tradable signal share the same no-news convention. Any score used in "
        "a portfolio is lagged by one equity session.",
    )
    h2(doc, "The sector index as an economic series")
    body(
        doc,
        "Figure A5 is a delayed description of the news climate. Treating it as a stand-alone timing rule would "
        "over-read the series. Context-weighted sector means are mildly positive over 2020 to 2023, but their "
        "ordering moves with the macro cycle. During the February to April 2020 COVID stress window, Energy is the "
        "weakest sector at about 0.005, whereas Technology and Real Estate sit closer to 0.05. That pattern fits "
        "an energy-demand shock and comparatively resilient digital and property-related language in the supplied "
        "corpus. It says nothing about realised sector returns.",
    )
    body(
        doc,
        "The ranking changes in 2022. Energy becomes the strongest sector on the index at roughly 0.061, while "
        "Materials and Communications sit near the bottom around 0.023 to 0.030. A year dominated by inflation, "
        "energy cash flows, and supply concerns makes that ordering economically plausible. Healthcare rises from "
        "about 0.036 in 2021 to 0.053 in 2023, and Technology recovers from 0.033 in 2022 to 0.051 in 2023.",
    )
    body(
        doc,
        "Coverage is not uniform. Technology and Consumer average more than five headlines on news-bearing "
        "ticker-days, whereas Materials and Real Estate are nearer two and have larger exact-zero raw-score "
        "shares (Figure A11). Context weighting discounts sparse observations so that a thinly covered sector is "
        "not treated as equally well measured. Communications stays in a relatively narrow 0.027 to 0.038 range "
        "across 2020 to 2023. Energy moves from 0.017 in 2020 to 0.061 in 2022 before easing to 0.041 in 2023, "
        "and Consumer remains roughly between 0.037 and 0.045. Neutral filling compresses the index level, so the "
        "useful information is mainly in relative movement across sectors and years. Figure A11 is the coverage "
        "check for whether a move reflects dense headline flow or a thin signal.",
    )
    h2(doc, "Continuity with Part A")
    body(
        doc,
        "Two design choices keep continuity with Part A. First, Attention Pulse remains a volume-confidence "
        "feature. It was descriptive in Part A and is not converted here into a contemporaneous trading trigger. "
        "Using only its elevated-volume confidence keeps the construction past-looking, without reversing the "
        "earlier claim that the pulse itself was not backtested alpha.",
    )
    body(
        doc,
        "Second, the phrase and lexicon layer is recorded in the audit material, not hidden inside a package "
        "label. The Week 9 FinVADER lexicons are vendored with Apache-2.0 notices, and Signal Harbour’s own "
        "terms, phrases, and masks are identified separately. Without that record, calling the model “augmented "
        "finVADER” would say very little about what actually changed.",
    )
    h2(doc, "Distributional model comparison")
    body(
        doc,
        "Table 4 and Figure A8 compare polarity shares across a 4,000-headline sample. A lower neutral share is "
        "only a distributional observation, not evidence of better prediction. I labelled the 120-headline "
        "worksheet myself. Signal Harbour buckets agree with those labels on 60.8% of the full sample "
        "(56.0% development, 68.9% holdout; n=120). The automated pseudo-label field is "
        "not treated as ground truth. Figure A9 looks at component effects through ablation on the Week 9 "
        "analyser and does not report labelled precision or recall. On the 800-headline ablation sample, "
        "removing neutral overrides has no effect. <i>no_neutral_overrides</i> matches <i>full_signal_harbour</i> "
        "because the retained masks rarely fire. The components that visibly move the neutral share are custom "
        "terms and phrases. Figure A10 reports active ticker counts. Figure A11 is the separate coverage-intensity "
        "exhibit.",
    )
    caption(
        doc,
        "Table 4. Model polarity shares on 4,000 headlines. The comparison is distributional only. "
        "Source: sentiment_model_comparison.csv; Figure A8.",
    )
    add_table(
        doc,
        ["Model", "Mean", "Positive", "Neutral", "Negative"],
        [
            ["Base VADER", "0.106", "37.5%", "49.6%", "12.9%"],
            ["Week 9 finVADER", "0.069", "26.6%", "62.1%", "11.3%"],
            ["Signal Harbour", "0.081", "29.8%", "58.7%", "11.6%"],
        ],
    )

    h2(doc, "Fusion result")
    body(
        doc,
        "The main fusion test applies the lagged context-weighted score to Equity Minimum Variance. Figure A6 "
        "shows that the two wealth paths are almost indistinguishable. The overlay is slightly worse. The "
        "zero-risk-free-rate Sharpe ratio falls from 0.4740 to 0.4711 and CAGR from 5.36% to 5.32%. Figure A7 "
        "and Table 6 extend the stress across tilt strengths and base portfolios. None of the tested "
        "specifications overtakes Equity Equal Weight. Table 5 then applies the same turnover-cost assumptions "
        "to the MinVar base and the fusion sleeve. Even modest costs preserve, and slightly widen, the "
        "disadvantage of the overlay. What remains usable is a continuous, look-ahead-safe monitoring pipeline, "
        "not a marketed alpha claim.",
    )
    caption(
        doc,
        "Table 5. Base versus fusion under equal turnover costs. TO is average annual turnover. "
        "Source: fusion_cost_sensitivity.csv.",
    )
    add_table(
        doc,
        ["Sleeve", "Cost", "Sharpe", "CAGR", "TO", "Terminal"],
        [
            ["Base", "0 bp", "0.474", "5.36%", "1.90", "1.169x"],
            ["Fusion", "0 bp", "0.471", "5.32%", "1.91", "1.167x"],
            ["Base", "5 bp", "0.467", "5.26%", "1.90", "1.165x"],
            ["Fusion", "5 bp", "0.464", "5.22%", "1.91", "1.164x"],
            ["Base", "10 bp", "0.459", "5.16%", "1.90", "1.162x"],
            ["Fusion", "10 bp", "0.456", "5.12%", "1.91", "1.161x"],
            ["Base", "25 bp", "0.437", "4.86%", "1.90", "1.152x"],
            ["Fusion", "25 bp", "0.434", "4.82%", "1.91", "1.151x"],
        ],
    )
    caption(
        doc,
        "Table 6. Fusion robustness across equity bases and tilt strengths at zero cost. "
        "Source: fusion_sensitivity.csv; Figure A7.",
    )
    add_table(
        doc,
        ["Base", "Tilt", "CAGR", "Sharpe", "MaxDD"],
        [
            ["EW", "0.15", "13.24%", "0.853", "-20.2%"],
            ["EW", "0.35", "13.17%", "0.849", "-20.2%"],
            ["EW", "0.60", "13.08%", "0.845", "-20.3%"],
            ["RP", "0.15", "10.39%", "0.753", "-18.4%"],
            ["RP", "0.35", "10.34%", "0.751", "-18.4%"],
            ["RP", "0.60", "10.29%", "0.748", "-18.5%"],
            ["MinVar", "0.15", "5.34%", "0.473", "-14.9%"],
            ["MinVar", "0.35", "5.32%", "0.471", "-14.9%"],
            ["MinVar", "0.60", "5.29%", "0.469", "-14.9%"],
        ],
    )
    body(
        doc,
        "The negative fusion result does not mean finance-aware sentiment has no research value. Table 4 and "
        "Figures A8 and A9 show that Week 9 finVADER and Signal Harbour redistribute polarity mass relative to "
        "social-media VADER, and the Explain view can still identify the terms and confidence weights behind a "
        "lagged score. In this 2021 to 2023 equity sample, a multiplicative portfolio tilt was not rewarded after "
        "the one-session lag, which sits with headline tone being partly contemporaneous with price information "
        "already present in the estimation window, and with the lag removing the most mechanical portion of any "
        "apparent edge. Repeating the test on Equal Weight and Risk Parity bases also shows the result is not an "
        "artefact of a damaged MinVar path. Thus, Recommendation 3 keeps sentiment as a monitoring and "
        "explanation layer, not as an independent return engine.",
    )

    h1(doc, "4. The Signal Harbour application")
    body(
        doc,
        "The application follows the analyst’s pre-committee workflow. Each tab answers a different question "
        "using precomputed research files, which keeps the free-tier deployment light and makes the displayed "
        "evidence reproducible.",
        indent=False,
    )
    body(
        doc,
        "<b>Compare.</b> Rank sleeves by zero-risk-free-rate Sharpe and terminal growth, then plot a shortlist "
        "of wealth paths from the research sample through December 2023.",
    )
    body(
        doc,
        "<b>Fact sheet.</b> Inspect one fund at a time: CAGR, volatility, Sharpe, drawdown, and the latest "
        "decision weights. A residual is stated whenever only the top holdings are displayed.",
    )
    body(
        doc,
        "<b>Allocate.</b> Combine completed fund return streams without re-optimising names, using intersecting "
        "equity sessions for equity-only blends and the native union calendar for blends containing crypto.",
    )
    body(
        doc,
        "<b>Sentiment and Explain.</b> Review the sector news climate, model comparisons, coverage diagnostics, "
        "and the prior-session fields that produced the lagged tradable score. Headlines are not re-scored live.",
    )
    h2(doc, "Calendar treatment in Allocate")
    body(
        doc,
        "Allocate uses a different calendar rule from the Combined-fund backtest because it mixes already-built "
        "return streams without re-optimising the underlying names. Any selection containing a Crypto fund "
        "uses the native union calendar, annualises at 365 days, and records zero for a closed equity leg. "
        "Equity-only blends stay on intersecting equity sessions and use 252-day annualisation. Due to that split, "
        "a 60/40 equity-crypto research allocation can retain crypto weekend returns without silently dropping "
        "Saturdays.",
    )
    body(
        doc,
        "A few product choices are worth stating because they affect what a user actually sees. The light theme "
        "is pinned so a dark-mode browser cannot place white headings on a pale custom background, Allocate does "
        "not imply that a crypto sleeve exists only on equity-open dates, and Explain displays the signal-date "
        "fields that entered the one-session lag so same-day headlines are not paired with yesterday’s tradable "
        "score. The fact sheet states the residual weight whenever the top-20 table is not the complete portfolio.",
    )
    body(
        doc,
        "For the advisory firm described in Part A, the interface maps onto a short meeting sequence. Begin with "
        "Combined Risk Parity and its drawdown profile (Recommendation 1; Sharpe 0.91). Combined Maximum Sharpe "
        "enters only when a client asks for more growth and accepts a deeper loss profile (Recommendation 2). "
        "Explain is for news-context colour, with Recommendation 3 already setting the boundary that sentiment is "
        "a monitoring layer, not a trading engine. If Energy’s stronger 2022 news-climate reading is "
        "questioned, the answer is in Section 3. Every series is labelled as an out-of-sample research path "
        "through December 2023, not as a live production feed.",
    )

    h1(doc, "5. Recommendations")
    body(
        doc,
        "The three recommendations below are taken from the verified tables and exhibits.",
        indent=False,
    )
    body(
        doc,
        "<b>1. Default mixed sleeve.</b> Use Combined Risk Parity as the primary balanced multi-asset comparison "
        "within this sample, with a zero-risk-free-rate Sharpe ratio of 0.91 and a maximum drawdown of -19.3% "
        "(Table 1; Figures A1, A2 and A4).",
    )
    body(
        doc,
        "<b>2. Satellite growth sleeve.</b> Keep Combined Maximum Sharpe as an optional higher-growth alternative "
        "rather than the default. Terminal wealth reaches 1.55x, but maximum drawdown deepens to -26.2% and the "
        "latest holdings are more concentrated (Tables 1 and 3; Figure A3).",
    )
    body(
        doc,
        "<b>3. Sentiment as monitoring, not alpha.</b> Retain the lagged context-weighted layer in the "
        "explainability workflow, but do not market it as a return engine. Relative to Equity MinVar, Sharpe "
        "moves from 0.4740 to 0.4711 and deteriorates further under equal cost assumptions (Figure A6; Tables 5 "
        "and 6).",
    )

    h1(doc, "6. Critical reflection and limits")
    body(
        doc,
        "Several choices make the study internally consistent, and they also narrow what can be claimed from it.",
        indent=False,
    )
    body(
        doc,
        "<b>Short live window.</b> Equity and Combined returns run from 5 January 2021 to December 2023, so the "
        "rankings are exposed to the COVID recovery, the 2022 inflation and rate shock, and the 2023 equity "
        "rebound. A longer sample could reorder Equal Weight and Risk Parity.",
    )
    body(
        doc,
        "<b>Headline results exclude costs.</b> Table 1 is a zero-cost laboratory, and Table 5 shows that even 5 "
        "to 25 basis points per unit of turnover reduce the fusion result. Applying the same grid to every fund "
        "would probably narrow the gap between higher-turnover Maximum Sharpe paths and lower-turnover Risk "
        "Parity.",
    )
    body(
        doc,
        "<b>Uneven news coverage.</b> Materials and Real Estate have thinner headline intensity and larger "
        "exact-zero raw-score shares than Technology or Consumer (Figure A11), so their sector readings require "
        "greater caution.",
    )
    body(
        doc,
        "<b>Labelled sentiment evidence is a small worksheet, not a precision claim.</b> I labelled all 120 "
        "headlines in <i>kaiyuan_review_label</i>. Agreement with Signal Harbour buckets is 60.8% on that sheet "
        "(headline_kaiyuan_label_agreement.csv). That is a polarity-match rate on n=120, not labelled precision "
        "or recall.",
    )
    body(
        doc,
        "<b>End-point holdings are not path explanations.</b> The December 2023 snapshots in Tables 2 and 3 "
        "describe concentration at the last decision date and cannot explain the complete 2021 to 2023 wealth "
        "path.",
    )
    h2(doc, "Product-level reflection")
    body(
        doc,
        "Building Signal Harbour on the actual Week 9 analyser was a stronger continuity choice than attaching "
        "a few finance terms to plain VADER. The report records the provenance of both the course lexicons and "
        "the project-specific additions. The application is explicit that it reads precomputed files and does "
        "not re-score headlines at runtime. Under those conditions, credibility comes from design quality, "
        "auditability, and restrained claims, not from the appearance of a live model.",
    )
    body(
        doc,
        "If I extended the project, I would make three changes. First, I would apply the turnover-cost grid "
        "from Table 5 to every fund family so that Recommendation 1 is tested under the same frictions. Second, "
        "I would expand the labelled worksheet beyond n=120 before treating polarity agreement as a "
        "model-selection statistic. Third, I would publish "
        "the public GitHub repository and Streamlit deployment, then resolve any remaining theme issues before "
        "submission. Those steps would improve validation and delivery without altering the central result that "
        "the sentiment overlay is economically flat to negative in this sample.",
    )
    h2(doc, "What out-of-sample means here")
    body(
        doc,
        "The fund weights are walk-forward and use only past returns, and the tradable sentiment input is delayed "
        "by one equity session. The research choices around lexicon retention, tilt strength, and the base sleeve "
        "receiving the overlay were nevertheless made with knowledge of the full study period, which is a normal "
        "limitation in a fixed-sample coursework project. The negative fusion result should be read as a stress "
        "test of one pre-specified overlay design, not as proof that every possible news rule failed under "
        "a sealed holdout. The development-holdout split in the headline worksheet offers partial discipline for "
        "labelling work, but it is not a complete untouched test set.",
    )
    body(
        doc,
        "Signal Harbour gives a small advisory desk a coherent menu of systematic funds, a lagged sector "
        "news-climate monitor linked to Part A, and an auditable explanation for why a name was tilted. It should "
        "not claim that the tilt paid in this sample, nor hide the calendar, cost, or coverage frictions that an "
        "investment committee would reasonably raise. Remaining packaging work is operational: keep the light "
        "theme pinned and verify that the live URL matches the version described here.",
    )

    doc.add_page_break()
    h1(doc, "Appendix A. Exhibits")
    body(
        doc,
        "Each exhibit is discussed in Sections 2 to 4. Captions restate the sample, source, and the claim the "
        "figure is meant to support.",
        indent=False,
    )
    exhibits = [
        (
            "growth_of_one_dollar.png",
            "Figure A1. Growth of $1 across funds, out-of-sample through December 2023. Combined Risk Parity is "
            "steadier than Combined Maximum Sharpe; Crypto Minimum Variance leads terminal wealth among the crypto "
            "sleeves. Source: fund_returns.csv.",
        ),
        (
            "combined_drawdowns.png",
            "Figure A2. Combined-fund drawdowns. Risk Parity has the shallower path, supporting Recommendation 1; "
            "Maximum Sharpe reaches the deeper trough expected of the higher-growth satellite. Wealth paths prepend "
            "1.0 so that the first loss is visible.",
        ),
        (
            "combined_weights_over_time.png",
            "Figure A3. Combined Minimum Variance and Risk Parity decision weights, with an Other residual so the "
            "stacked areas sum to one. Maximum Sharpe concentration is summarised in Table 3 rather than in this panel.",
        ),
        (
            "fund_sharpe_barplot.png",
            "Figure A4. Zero-risk-free-rate Sharpe rankings. Combined Risk Parity leads the mixed sleeves; Equity "
            "Equal Weight leads equities; Crypto Minimum Variance leads crypto. Annualisation is 252 or 365 days by sleeve.",
        ),
        (
            "sector_sentiment_index.png",
            "Figure A5. Context-weighted sector sentiment, shown as a 21-session average. The series uses the complete "
            "ticker-day grid, sets no-news observations to neutral, and is lagged before any trading use.",
        ),
        (
            "fusion_before_after.png",
            "Figure A6. Equity Minimum Variance against the lagged context-weighted sentiment tilt. The overlay does "
            "not improve the corrected baseline, which supports Recommendation 3.",
        ),
        (
            "fusion_sensitivity.png",
            "Figure A7. Fusion sensitivity by tilt strength and equity base at zero cost. None of the tested "
            "specifications overtakes Equity Equal Weight (Table 6).",
        ),
        (
            "sentiment_model_comparison.png",
            "Figure A8. Polarity shares for Base VADER, Week 9 finVADER, and Signal Harbour on 4,000 headlines "
            "(Table 4). The comparison is distributional only.",
        ),
        (
            "innovation_ablation.png",
            "Figure A9. Ablation of Signal Harbour components on the Week 9 finVADER base analyser, with Week 9-only "
            "and plain-VADER references.",
        ),
        (
            "context_weighted_sector_sentiment.png",
            "Figure A10. Average active tickers per sector-day under the complete-grid context-weighted index. This "
            "is the coverage companion to Figure A5.",
        ),
        (
            "sentiment_coverage_neutrality.png",
            "Figure A11. Sector coverage intensity: mean headlines on news-bearing ticker-days (bars) and the share "
            "of exact-zero raw scores (line). This exhibit is distinct from the polarity comparison in Figure A8.",
        ),
    ]
    for i, (fname, cap) in enumerate(exhibits):
        if i > 0:
            doc.add_page_break()
        add_figure(doc, fname, cap)

    doc.add_page_break()
    h1(doc, "Appendix B. Core equations and references")
    h2(doc, "Minimum variance")
    eq(doc, "minimise  w^T SigmaAnn w    subject to    1^T w = 1    and    0 <= w(i) <= wMax")
    eqnote(
        doc,
        "SigmaAnn = ppy * Cov(R), where ppy is 252 on the equity calendar or 365 on the native crypto calendar.",
    )
    h2(doc, "Maximum Sharpe")
    eq(doc, "maximise  (w^T muAnn) / sqrt(w^T SigmaAnn w)")
    eqnote(doc, "The risk-free rate is zero; the long-only and name-cap constraints are unchanged.")
    h2(doc, "Reported performance")
    eq(doc, "Sharpe = (mean(r) * ppy) / (std(r) * sqrt(ppy))")
    eq(doc, "CAGR = WT raised to the power (ppy/T), minus 1")
    h2(doc, "Context-weighted score")
    eq(doc, "sCw(i, t) = s(i, t) * coverageConf(i, t) * attentionConf(i, t)")
    eq(doc, "Tradable score = sCw(i, t minus 1).")
    h2(doc, "Sentiment tilt")
    eqnote(
        doc,
        "wTilt is proportional to w * exp(lambda * (score minus mean(score))). "
        "The tilted vector is then projected back onto the capped simplex.",
    )
    h2(doc, "References")
    ref(
        doc,
        "Hutto, C. J., and Gilbert, E. (2014). VADER: A parsimonious rule-based model for sentiment analysis of social media text. <i>ICWSM</i>.",
    )
    ref(
        doc,
        "Korab, P. FinVADER lexicons (SentiBigNomics / Henry path used in FINS3645 Week 9). Apache 2.0; vendored under src/vendor/course_finvader.",
    )
    ref(
        doc,
        "Loughran, T., and McDonald, B. (2011). When is a liability not a liability? <i>The Journal of Finance</i>, 66(1), 35 to 65.",
    )
    ref(
        doc,
        "Markowitz, H. (1952). Portfolio selection. <i>The Journal of Finance</i>, 7(1), 77 to 91.",
    )
    ref(
        doc,
        "FINS3645 Project 2026 brief and data access helper; Kaiyuan Lan, Project A Signal Harbour / Attention Pulse (2026).",
    )

    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    print(build())
